import json
import math
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "rows-36-39.json"
OUTPUT = ROOT / "คำตอบนักเรียน-แถว-36-39.docx"


def set_run_font(run, name="Arial", size=11, bold=False, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="DADCE0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is not None:
        tbl_w.set(qn("w:w"), "9360")
        tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def excel_serial_to_text(value):
    if not isinstance(value, (int, float)) or isinstance(value, bool):
        return None
    if not math.isfinite(value) or value < 1 or value > 100000:
        return None
    base = datetime(1899, 12, 30)
    dt = base + timedelta(days=float(value))
    return dt.strftime("%Y-%m-%d %H:%M:%S")


def value_text(column, value):
    if value is None:
        return ""
    if column == "B":
        date_text = excel_serial_to_text(value)
        if date_text:
            return date_text
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    return str(value)


data = json.loads(SOURCE.read_text(encoding="utf-8"))
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title_run = title.add_run("คำตอบนักเรียน แถว 36–39")
set_run_font(title_run, size=26)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
sub_run = subtitle.add_run("คัดลอกข้อมูลทั้งแถวจากชีต “คำตอบนักเรียน” ในไฟล์ SDS-data (1).xlsx")
set_run_font(sub_run, size=10, color="555555")

for row_index, row in enumerate(data["rows"]):
    if row_index:
        doc.add_page_break()
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(f"แถวที่ {row['row']}")
    set_run_font(run, size=20)

    name_value = value_text("C", row["values"][2])
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run(f"นักเรียน: {name_value or '-'}")
    set_run_font(meta_run, size=11, bold=True)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2088, 7272])
    table.rows[0].cells[0].text = "คอลัมน์"
    table.rows[0].cells[1].text = "ค่า"
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True)

    for header, value in zip(data["headers"], row["values"]):
        cells = table.add_row().cells
        cells[0].text = f"{header['column']} — {header['header'] or '(ไม่มีชื่อคอลัมน์)'}"
        cells[1].text = value_text(header["column"], value)
        for index, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=9 if index == 0 else 9.5)

doc.save(OUTPUT)
print(OUTPUT)
